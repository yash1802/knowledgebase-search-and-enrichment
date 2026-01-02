import re
from pathlib import Path
import PyPDF2
from docx import Document
from config.settings import (
    USE_PAGE_LEVEL_CHUNKING,
    MAX_PAGE_SIZE,
    MIN_PAGE_SIZE,
    SEMANTIC_CHUNK_SIZE,
    SEMANTIC_CHUNK_OVERLAP,
    MANUAL_INFO_FILENAME
)


class DocumentProcessor:
    """Process documents: extract text, clean, and chunk."""

    def __init__(self, chunk_size = 1000, chunk_overlap = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.use_page_level = USE_PAGE_LEVEL_CHUNKING
        self.semantic_chunk_size = SEMANTIC_CHUNK_SIZE
        self.semantic_chunk_overlap = SEMANTIC_CHUNK_OVERLAP
        self.max_page_size = MAX_PAGE_SIZE
        self.min_page_size = MIN_PAGE_SIZE

    def extract_text(self, file_path):
        file_path_obj = Path(file_path)
        extension = file_path_obj.suffix.lower()

        if extension == ".pdf":
            return self._extract_from_pdf(file_path)
        elif extension == ".txt" or extension == ".md":
            return self._extract_from_text(file_path)
        elif extension == ".docx":
            return self._extract_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")

    def _extract_from_pdf(self, file_path):
        text = ""
        try:
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
        except Exception as e:
            raise ValueError(f"Error reading PDF: {str(e)}")
        return text.strip()

    def _extract_from_text(self, file_path):
        with open(file_path, "r", encoding="utf-8") as file:
            return file.read()

    def _extract_from_docx(self, file_path):
        try:
            doc = Document(file_path)
            return "\n".join([paragraph.text for paragraph in doc.paragraphs])
        except Exception as e:
            raise ValueError(f"Error reading DOCX: {str(e)}")

    def clean_text(self, text):
        # Remove whitespace and special characters but keep punctuation
        text = re.sub(r"\s+", " ", text)
        text = re.sub(r"[^\w\s.,;:!?\-()\[\]{}'\"\/]", " ", text)
        return text.strip()

    def chunk_text(self, text):
        if len(text) <= self.chunk_size:
            return [text]

        chunks = []
        start = 0

        while start < len(text):
            end = start + self.chunk_size
            chunk = text[start:end]

            # Try to break at sentence boundary
            if end < len(text):
                last_period = chunk.rfind(".")
                last_newline = chunk.rfind("\n")
                break_point = max(last_period, last_newline)

                if break_point > self.chunk_size * 0.5:
                    chunk = chunk[:break_point + 1]
                    end = start + break_point + 1

            chunks.append(chunk.strip())

            # Move start position with overlap
            start = end - self.chunk_overlap
            if start >= len(text):
                break

        return chunks

    def _chunk_pdf_by_pages(self, file_path):
        chunks = []
        try:
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    
                    if page_text:
                        cleaned_text = self.clean_text(page_text)
                        
                        if len(cleaned_text) > self.max_page_size:
                            page_chunks = self._split_large_page(cleaned_text)
                            chunks.extend(page_chunks)
                        elif len(cleaned_text) >= self.min_page_size:
                            chunks.append(cleaned_text)
                        else:
                            print(f"  Page {page_num + 1}: FILTERED OUT (too small, min={self.min_page_size})")
                        
        except Exception as e:
            raise ValueError(f"Error chunking PDF by pages: {str(e)}")
        
        
        validated = self._validate_chunks(chunks)
        return validated

    def _chunk_docx_by_pages(self, file_path):
        try:
            doc = Document(file_path)
            chunks = []
            current_chunk = []
            current_size = 0
            
            for paragraph in doc.paragraphs:
                para_text = self.clean_text(paragraph.text)
                para_size = len(para_text)
                
                # If adding this paragraph exceeds max page size, start new chunk
                if current_size + para_size > self.max_page_size and current_chunk:
                    chunk_text = " ".join(current_chunk)
                    if len(chunk_text) >= self.min_page_size:
                        chunks.append(chunk_text)
                    current_chunk = [para_text] if para_text else []
                    current_size = para_size
                else:
                    if para_text:
                        current_chunk.append(para_text)
                        current_size += para_size
            
            if current_chunk:
                chunk_text = " ".join(current_chunk)
                if len(chunk_text) >= self.min_page_size:
                    chunks.append(chunk_text)
            
            return self._validate_chunks(chunks)
        except Exception as e:
            raise ValueError(f"Error chunking DOCX by pages: {str(e)}")

    def _chunk_markdown(self, text):
        """
        Chunk markdown by sections (headers).
        """
        # Split by markdown headers (# ## ### etc)
        header_pattern = r'^#{1,6}\s+.+$'
        lines = text.split('\n')
        
        chunks = []
        current_chunk = []
        current_size = 0
        
        for line in lines:
            line_size = len(line)
            
            # Check if this is a header
            if re.match(header_pattern, line.strip()):
                # Save previous chunk if it exists
                if current_chunk:
                    chunk_text = self.clean_text('\n'.join(current_chunk))
                    if len(chunk_text) >= self.min_page_size:
                        chunks.append(chunk_text)
                
                # Start new chunk with this header
                current_chunk = [line]
                current_size = line_size
            else:
                # adding this line exceeds semantic chunk size, save and start new
                if current_size + line_size > self.semantic_chunk_size and current_chunk:
                    chunk_text = self.clean_text('\n'.join(current_chunk))
                    if len(chunk_text) >= self.min_page_size:
                        chunks.append(chunk_text)
                    current_chunk = [line]
                    current_size = line_size
                else:
                    current_chunk.append(line)
                    current_size += line_size
        
        if current_chunk:
            chunk_text = self.clean_text('\n'.join(current_chunk))
            if len(chunk_text) >= self.min_page_size:
                chunks.append(chunk_text)
        
        return self._validate_chunks(chunks)

    def _chunk_by_paragraphs(self, text):
        """
        Chunk text by paragraphs with overlap.
        """
        
        paragraphs = re.split(r'\n\s*\n', text)
        
        chunks = []
        current_chunk = []
        current_size = 0
        
        for para in paragraphs:
            para_clean = self.clean_text(para)
            para_size = len(para_clean)
            
            # Adding this paragraph exceeds semantic chunk size, save and start new chunk
            if current_size + para_size > self.semantic_chunk_size and current_chunk:
                chunk_text = ' '.join(current_chunk)
                if len(chunk_text) >= self.min_page_size:
                    chunks.append(chunk_text)
                
                # Start new chunk with overlap (keep last paragraph)
                if len(current_chunk) > 1:
                    current_chunk = [current_chunk[-1], para_clean]
                    current_size = len(current_chunk[-2]) + para_size
                else:
                    current_chunk = [para_clean]
                    current_size = para_size
            else:
                if para_clean:
                    current_chunk.append(para_clean)
                    current_size += para_size
        
        # Add remaining chunk
        if current_chunk:
            chunk_text = ' '.join(current_chunk)
            if len(chunk_text) >= self.min_page_size:
                chunks.append(chunk_text)
        
        return self._validate_chunks(chunks)

    def _split_large_page(self, text):
        """
        Split a large page into smaller chunks.
        """
        # Use the standard chunking with semantic chunk size
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.semantic_chunk_size
            chunk = text[start:end]
            
            # Try to break at sentence boundary
            if end < len(text):
                last_period = chunk.rfind(".")
                last_newline = chunk.rfind("\n")
                break_point = max(last_period, last_newline)
                
                if break_point > self.semantic_chunk_size * 0.5:
                    chunk = chunk[:break_point + 1]
                    end = start + break_point + 1
            
            if chunk.strip():
                chunks.append(chunk.strip())
            
            start = end - self.semantic_chunk_overlap
            if start >= len(text):
                break
        
        return chunks


    def _validate_chunks(self, chunks):
        """
        Ensure chunks meet size requirements.
        """
        validated = []
        filtered_count = 0
        split_count = 0
        
        for i, chunk in enumerate(chunks):
            chunk_len = len(chunk)
            # Filter out chunks that are too small
            if chunk_len >= self.min_page_size:
                # Truncate chunks that are too large
                if chunk_len > self.max_page_size:
                    # Split into smaller chunks
                    split_chunks = self._split_large_page(chunk)
                    validated.extend(split_chunks)
                    split_count += 1
                else:
                    validated.append(chunk)
            else:
                filtered_count += 1
        
        
        return validated

    def _chunk_manual_information_file(self, text):
        """
        Parse manual_information.txt into chunks.
        Strips timestamps, only returns content.
        """
        chunks = []
        entries = text.split("\n\n")
        
        for entry in entries:
            entry = entry.strip()
            if entry:
                lines = entry.split("\n")
                # First line should be timestamp [YYYY-MM-DD HH:MM:SS]
                if len(lines) >= 2 and lines[0].startswith("[") and lines[0].endswith("]"):
                    # Extract content (skip timestamp line)
                    content = "\n".join(lines[1:]).strip()
                    if content:
                        chunks.append(content)
                else:
                    # Malformed entry, include as-is
                    chunks.append(entry)
        
        return chunks

    def process_document(self, file_path):
        """
        Process a document: extract, clean, and chunk using file-type dependent strategy.
        """
        file_type = Path(file_path).suffix.lower()
        filename = Path(file_path).name
        
        # Special handling for manual_information.txt
        if filename == MANUAL_INFO_FILENAME:
            text = self.extract_text(file_path)
            cleaned_text = self.clean_text(text)
            chunks = self._chunk_manual_information_file(cleaned_text)
        
        # Choose chunking strategy based on file type
        elif file_type == ".pdf" and self.use_page_level:
            chunks = self._chunk_pdf_by_pages(file_path)
            # Extract full text for metadata
            text = self.extract_text(file_path)
            cleaned_text = self.clean_text(text)
        elif file_type == ".docx" and self.use_page_level:
            chunks = self._chunk_docx_by_pages(file_path)
            # Extract full text for metadata
            text = self.extract_text(file_path)
            cleaned_text = self.clean_text(text)
        elif file_type == ".md":
            text = self.extract_text(file_path)
            cleaned_text = self.clean_text(text)
            chunks = self._chunk_markdown(cleaned_text)
        elif file_type == ".txt":
            text = self.extract_text(file_path)
            cleaned_text = self.clean_text(text)
            chunks = self._chunk_by_paragraphs(cleaned_text)
        else:
            # Fallback to standard chunking
            text = self.extract_text(file_path)
            cleaned_text = self.clean_text(text)
            chunks = self.chunk_text(cleaned_text)


        return {
            "file_path": file_path,
            "filename": Path(file_path).name,
            "text": cleaned_text,
            "chunks": chunks,
            "num_chunks": len(chunks)
        }

